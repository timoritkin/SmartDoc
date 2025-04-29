import os
import shutil
import subprocess
import sys
import time
import tkinter as tk
from datetime import datetime
from tkinter import messagebox, ttk
from docx import Document
from docxtpl import DocxTemplate
import db_setup as db
import customtkinter as ctk
from PIL import Image
from tkcalendar import DateEntry
from customtkinter import CTkImage
from pathlib import Path

hebrew_font = ("Arial", 16, "bold")
padX_size = (30, 30)
padY_size = (0, 20)
sticky_label = "w"
sticky_entry = "e"
borders_widgets = (30, 20)
color1 = "#176B87"
color2 = "#64CCC5"

rest_image = Image.open("images/icons8-restart-50.png")
rest_icon = ctk.CTkImage(dark_image=rest_image, size=(20, 20))  # Adjust size as needed

new_user_image = Image.open("images/icons8-add-user-male-50.png")
new_user_icon = ctk.CTkImage(dark_image=new_user_image, size=(20, 20))  # Adjust size as needed

search_user_image = Image.open("images/icons8-find-user-male-50.png")
search_user_icon = ctk.CTkImage(dark_image=search_user_image, size=(20, 20))  # Adjust size as needed

search_form_image = Image.open("images/icons8-search-property-50.png")
search_form_icon = ctk.CTkImage(dark_image=search_form_image, size=(20, 20))  # Adjust size as needed

# Get the user's Local AppData folder (C:\Users\YourUsername\AppData\Local)
local_appdata_path = Path(os.getenv('LOCALAPPDATA') or "") / 'SmartDoc'

# Create a subfolder for your app inside Local AppData
db_folder = local_appdata_path / 'Database'
db_folder.mkdir(parents=True, exist_ok=True)  # Ensure the folder exists

# Full path to the database file
db_path = db_folder / 'patients.db'

# Define base and patient-specific folders inside Local AppData
patients_base_folder = local_appdata_path / 'My Patients'
patients_base_folder.mkdir(parents=True, exist_ok=True)


def update_text_in_docx(old_data, new_data):
    text_to_update = {}
    labels = ["טלפון", "גיל", "שם פרטי", "שם משפחה", "תעודת זהות"]

    change_file_name = False  # only specific filed will force to change folder and file names

    new_patient_data = {}
    for key in new_data:
        if key == "גיל":
            new_patient_data[key] = new_data[key].get_date().strftime('%d/%m/%Y')
            new_patient_data[key] = calculate_age(new_patient_data[key])
            print(new_patient_data[key])
        else:
            new_patient_data[key] = new_data[key].get()

    for i in range(len(labels)):
        print(old_data[i])
        old_value = str(old_data[i])
        label = labels[i]
        new_value = str(new_patient_data[label])
        if old_value != new_value:
            if i == 2 or i == 3 or i == 4:
                change_file_name = True
            text_to_update[old_value] = new_value

    # Get the path to the existing .docx file
    print(old_data[4])
    file_path_tuple = db.get_patient_docx_path(new_data["תעודת זהות"].get(), db_path)
    print(file_path_tuple)
    file_path = Path(file_path_tuple[0])  # Use Path directly

    if not file_path.exists():
        print(" File not found:", file_path)
        return
    if file_path.suffix.lower() != ".docx":
        print(" Not a .docx file:", file_path)
        return

    try:
        doc = Document(str(file_path))
    except Exception as e:
        print(" Failed to load document:", e)
        return

    # Perform replacements in tables
    for old_text, new_text in text_to_update.items():
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)

    # Save the updated document to the same path (overwrite it)
    doc.save(str(file_path))
    print(" Updated and saved document at:", file_path)

    if change_file_name:

        # Create patient folder structure
        new_patient_folder = patients_base_folder / f"{new_data['שם פרטי'].get()}_{new_data['שם משפחה'].get()}_{new_data['תעודת זהות'].get()}"
        create_directory(new_patient_folder)  # Using the create_directory function consistently
        print(new_patient_folder)
        # Create filename with date included
        new_file_name = f"{new_data['שם פרטי'].get()}_{new_data['שם משפחה'].get()}_{new_data['תעודת זהות'].get()}.docx"
        new_file_path = new_patient_folder / new_file_name
        print(new_file_path)
        source_path = Path(file_path)
        destination_path = Path(new_file_path)

        if source_path.exists():
            shutil.copy(source_path, destination_path)
            print(f"Copied {source_path} to {destination_path}")
            db.update_docx_path(new_data['תעודת זהות'].get(), destination_path, db_path)

            # delete process of old path
            # Delete the file first
            if file_path.exists() and file_path.is_file():
                file_path.unlink()
                print(f"Deleted file: {file_path}")

                # Now check if the parent folder is empty
                parent_folder = file_path.parent
                if not any(parent_folder.iterdir()):  # Folder is empty
                    parent_folder.rmdir()
                    print(f"Deleted empty folder: {parent_folder}")
                    messagebox.showinfo("עודכן", "הנתונים עודכנו בהצלחה")

            else:
                print(f"File does not exist: {file_path}")
                messagebox.showerror("שגיאה", f"File does not exist: {file_path}")

        else:
            print(f"File does not exist: {source_path}")
            messagebox.showerror("שגיאה", f"File does not exist: {source_path}")

    else:
        messagebox.showinfo("עודכן", "הנתונים עודכנו בהצלחה")


def update_age_of_patient_in_docx(old_age, new_age, path):
    file = Path(path)
    if not file.exists():
        print(" File not found:", file)
        return
    if file.suffix.lower() != ".docx":
        print(" Not a .docx file:", file)
        return

    try:
        doc = Document(str(file))
    except Exception as e:
        print(" Failed to load document:", e)
        return

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if old_age in run.text:
                            run.text = run.text.replace(old_age, new_age)

    # Save the updated document to the same path (overwrite it)
    doc.save(str(file))
    print(" Updated and saved document at:", file)


def sort_treeview_column(treeview, column, reverse):
    """
    Sort the Treeview by the specified column.

    :param treeview: The Treeview widget.
    :param column: The column to sort by.
    :param reverse: Boolean value to indicate whether to reverse the order.
    """
    # Get the data from the treeview and sort it by the column
    data = [(treeview.set(child, column), child) for child in treeview.get_children('')]
    data.sort(key=lambda x: x[0], reverse=reverse)

    # Re-insert the sorted data back into the treeview
    for index, item in enumerate(data):
        treeview.move(item[1], '', index)

    # Toggle the reverse order for the next sort
    return not reverse


def on_column_click(treeview, column, sort_directions):
    """
    Handle a column click for sorting.

    :param treeview: The Treeview widget.
    :param column: The column that was clicked.
    :param sort_directions: A dictionary storing the sort direction for each column.
    """
    # Get the current sorting direction for the clicked column
    reverse = sort_directions.get(column, False)  # Default to False (ascending) if not set

    # Toggle the sorting order (ascending <-> descending)
    sort_directions[column] = not reverse  # Toggle the value for the column

    # Sort the treeview based on the column and direction
    sort_treeview_column(treeview, column, reverse)


def resource_path(relative_path):
    """ Get absolute path to resource, works for development and for cx_Freeze"""
    if getattr(sys, 'frozen', False):  # Check if the app is frozen (packaged)
        # cx_Freeze uses sys.executable for the base path
        base_path = os.path.dirname(sys.executable)
    else:
        # Development mode
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)


# Get the selected item
def open_word_document(event):
    # Get the selected item
    selected_item = event.widget.selection()
    if not selected_item:
        return
    p_id = event.widget.item(selected_item, 'values')[5]
    p_current_age = event.widget.item(selected_item, 'values')[2]  # maybe the age changed

    # Retrieve file information
    visit_date = event.widget.item(selected_item, 'values')[0]
    path = db.get_docx_path(p_id, visit_date, db_path)
    p_birthdate = db.get_patient_birthdate(p_id, db_path)
    p_new_age = calculate_age(p_birthdate[0])
    print(p_current_age)
    print(p_new_age)
    print(path)

    # Get the document path from the database
    if path:
        path = resource_path(path)
    # Resolve the full path for bundled environments
    if path and os.path.exists(path):
        try:
            if str(p_new_age) != p_current_age:
                print("there is different age for the patient")
                update_age_of_patient_in_docx(p_current_age, p_new_age, path)
            else:
                print("there is no difference")

            # Use the default application to open the file
            if os.name == 'nt':  # Windows
                os.startfile(path)
            elif os.name == 'posix':  # macOS and Linux
                subprocess.run(['open', path], check=True)
            else:
                messagebox.showerror("Error", "Unsupported operating system")
        except Exception as e:
            messagebox.showerror("Error", f"Error opening file: {e}")
    else:
        messagebox.showwarning("Warning", "הקובץ לא נמצא")


def create_directory(path):
    """Ensure a directory exists."""
    path.mkdir(parents=True, exist_ok=True)


def open_file(file_path):
    """Open a file based on the operating system."""
    if sys.platform == "win32":  # For Windows
        os.startfile(file_path)
    elif sys.platform == "darwin":  # For macOS
        subprocess.run(["open", file_path])
    else:  # For Linux
        subprocess.run(["xdg-open", file_path])


def create_docx(f_name, l_name, id_num, age, phone):
    """
    Create a customized document from a template for a patient.

    Args:
        f_name (str): Patient's first name
        l_name (str): Patient's last name
        id_num (str): Patient's ID number
        age (str/int): Patient's age
        phone (str): Patient's phone number

    Returns:
        str: Path to the created document
    """
    # Load the template using resource_path
    template_path = resource_path('template/Clalit mushlam template.docx')
    doc = DocxTemplate(template_path)

    # Format phone number - add leading zero if needed
    phone = str(phone)
    if not phone.startswith("0") and len(phone) == 9:
        phone = "0" + phone

    # Create patient folder structure
    patient_folder = patients_base_folder / f"{f_name}_{l_name}_{id_num}"
    create_directory(patient_folder)  # Using the create_directory function consistently

    # Create filename with date included
    file_name = f"{f_name}_{l_name}_{id_num}.docx"
    file_path = patient_folder / file_name

    # Prepare context for the document
    context = {'f_name': f_name, 'l_name': l_name, 'id': id_num, 'age': age, 'phone': phone}

    # Render and save the document
    doc.render(context)
    doc.save(file_path)

    # Open the document automatically
    open_file(file_path)

    # Return the path as a string
    return str(file_path)


def adjust_data(tree):
    selected_item = tree.selection()  # Get selected row
    if not selected_item:  # If no row is selected, exit
        return

    item_data = tree.item(selected_item[0])["values"]  # Extract data
    patient_id = item_data[4]
    print("the id is", patient_id)
    # Create a pop-up window for editing
    popup = tk.Toplevel()
    popup.title("עריכה")
    popup.resizable(False, False)

    popup_frame = ctk.CTkFrame(popup, fg_color=color1)
    popup_frame.pack(padx=10, pady=10)

    # Form Labels and Entry Fields
    labels = ["טלפון", "גיל", "שם פרטי", "שם משפחה", "תעודת זהות"]
    fields = {}

    for i, label_text in enumerate(labels):
        label = ctk.CTkLabel(popup_frame, text=label_text, font=hebrew_font)
        label.grid(row=i, column=1, padx=10, pady=5, sticky="e")

        if label_text == "טלפון":
            # add the zero to the phone number
            phone = str(item_data[i])
            if not phone.startswith("0") and len(phone) == 9:
                phone = "0" + phone
                item_data[i] = phone
            print(phone)

            entry_var = ctk.StringVar(value=phone)  # Pre-fill with data
            entry = ctk.CTkEntry(popup_frame, textvariable=entry_var, width=250)
            entry.grid(row=i, column=0, padx=10, pady=5, sticky="w")
            fields[label_text] = entry_var

        elif label_text == "גיל":
            patient_birthdate = db.get_patient_birthdate(item_data[4], db_path)
            calendar = DateEntry(
                popup_frame,
                date_pattern='dd/mm/yyyy',
                width=24,
                background="darkblue",
                foreground="white",
                font=("Arial", 13),
                state="normal"
            )
            calendar.set_date(patient_birthdate[0])
            calendar.grid(row=i, column=0, padx=10, pady=5, sticky="w")
            fields[label_text] = calendar  # Store the widget (not StringVar!)
        else:
            entry_var = ctk.StringVar(value=item_data[i])  # Pre-fill with data
            entry = ctk.CTkEntry(popup_frame, textvariable=entry_var, width=250)
            entry.grid(row=i, column=0, padx=10, pady=5, sticky="w")
            fields[label_text] = entry_var

    def update_everything():
        db.update_patient_record(fields, patient_id, db_path)

        update_text_in_docx(item_data, fields)
        popup.destroy()

    # Submit Button to Save Changes
    create_button = ctk.CTkButton(
        popup_frame,
        text="נתונים עדכן",
        width=250,
        height=40,
        command=update_everything
    )
    create_button.grid(row=len(labels), column=0, columnspan=2, pady=(10, 20))


# user will be asked if they want to create new visit as a result new doc will be created
def create_new_visit(event, tree):
    # Get the selected row (item) that was double-clicked
    selected_item = tree.selection()

    # if selected_item:  # If an item is selected
    #     # Get the data (values) of the selected row
    #     item_data = tree.item(selected_item[0])["values"]
    #     # Get the current date in the desired format (e.g., dd-mm-yyyy)
    #     current_date = datetime.now().strftime('%d-%m-%Y')  # Use hyphens instead of slashes
    #
    #     # Create a new window (pop-up)
    #     popup = tk.Toplevel()
    #     popup.title("ביקור יצירת")
    #     # Prevent the window from being resized
    #     popup.resizable(False, False)
    #     popup_frame = ctk.CTkFrame(popup, fg_color=color1)  # Use ctk.CTkFrame directly
    #
    #     print(item_data[0])
    #     print(item_data[1])
    #     print(item_data[2])
    #     print(item_data[3])
    #     print(item_data[4])
    #
    #     # Last Name
    #     new_record_label = ctk.CTkLabel(
    #         popup_frame,
    #         text=f"? האם ליצור עבור המטופל {item_data[2]} {item_data[3]} ביקור חדש ",
    #         font=hebrew_font,
    #         anchor="e"
    #     )
    #     new_record_label.grid(row=0, column=0, columnspan=2, padx=padX_size, pady=padY_size, sticky='nswe')
    #
    #     confirm_button = ctk.CTkButton(popup_frame,
    #                                    text="אישור",
    #                                    width=100,
    #                                    command=lambda: ((docx_path := create_docx(item_data[2], item_data[3],
    #                                                                               item_data[4], item_data[1],
    #                                                                               item_data[0])),
    #                                                     db.insert_visit_record(item_data[4], current_date, docx_path,
    #                                                                            db_path),
    #                                                     popup.destroy()))
    #
    #     confirm_button.grid(row=1, column=1, sticky='we', padx=10, pady=10)
    #
    #     denied_button = ctk.CTkButton(popup_frame,
    #                                   text="ביטול",
    #                                   width=100,
    #                                   fg_color="red",
    #                                   hover_color="#AF1740",
    #                                   command=popup.destroy)
    #     denied_button.grid(row=1, column=0, sticky='we', padx=10, pady=10)
    #     popup_frame.grid(row=0, column=0, sticky="nsew")
    #


def load_visit_data(self):
    # Clear existing Treeview contents before inserting new data
    for item in self.visit_treeview.get_children():
        self.visit_treeview.delete(item)

    # Fetch data and populate the Treeview
    rows = db.fetch_visit_data(db_path)

    print(rows)

    for row in rows:
        birthdate_str = row[2]  # Example: row[2] is the birthdate column in 'dd/mm/yyyy' format
        age = calculate_age(birthdate_str)
        row_with_replaced_age = list(row)  # Convert the tuple to a list
        row_with_replaced_age[2] = age  # Assuming the Age column is at index 2
        self.visit_treeview.insert("", tk.END, values=row_with_replaced_age)


def load_patient_data(self):
    # Clear existing Treeview contents before inserting new data
    for item in self.patients_treeview.get_children():
        self.patients_treeview.delete(item)

    # Fetch data and populate the Treeview
    rows = db.fetch_patient_data(db_path)

    for row in rows:
        birthdate_str = row[1]  # Example: row[2] is the birthdate column in 'dd/mm/yyyy' format
        age = calculate_age(birthdate_str)
        row_with_replaced_age = list(row)  # Convert the tuple to a list
        row_with_replaced_age[1] = age  # Assuming the Age column is at index 2
        self.patients_treeview.insert("", tk.END, values=row_with_replaced_age)


def calculate_age(birthdate_str):
    try:
        # Parse the birthdate string
        birthdate = datetime.strptime(birthdate_str, '%d/%m/%Y')
    except ValueError:
        return None

    # Calculate the current age
    current_date = datetime.today()
    age = current_date.year - birthdate.year

    # Adjust for birthday not yet occurring this year
    if current_date.month < birthdate.month or (
            current_date.month == birthdate.month and current_date.day < birthdate.day):
        age -= 1

    return age


class PatientForm:

    def __init__(self, root):
        self.root = root
        self.root.title("SmartDoc")
        self.root.geometry("900x700")
        self.root.iconbitmap("logo/logo_icon.ico")  # Provide the path to your .ico file
        self.root.configure(bg=color1)  # Use a color name or hex code
        # Configure the grid layout for the window
        self.root.grid_columnconfigure(0, weight=1)  # Main frame will expand
        self.root.grid_columnconfigure(1, weight=0)  # Options frame stays fixed
        self.root.grid_rowconfigure(0, weight=1)  # Main frame will expand vertically
        # Main frame (left side)
        self.main_frame = ctk.CTkFrame(self.root, fg_color=color1)  # Use ctk.CTkFrame directly
        self.main_frame.grid(row=0, column=0, sticky="nsew")

        # Options frame (right side)
        self.options_frame = ctk.CTkFrame(self.root, fg_color=color2)  # Use ctk.CTkFrame directly
        self.options_frame.grid(row=0, column=1, sticky="ns")  # Stick to top and bottom
        # Load an rest_image using Pillow
        image = Image.open("logo/SmartDocLogo.png")
        ctk_image = CTkImage(light_image=image, size=(200, 100))
        self.error_is_raised = True
        self.logo_label = ctk.CTkLabel(
            self.options_frame,
            image=ctk_image,
            text=""  # Set text to an empty string to only show the rest_image
        )
        self.logo_label.pack(pady=(0, 150))
        # Adding buttons to options_frame
        self.new_form_button = ctk.CTkButton(self.options_frame,
                                             image=new_user_icon,
                                             text="חדש מטופל",
                                             width=200,
                                             height=40,
                                             command=self.show_new_form)
        self.new_form_button.pack(pady=10)
        self.search_visit_button = ctk.CTkButton(self.options_frame,
                                                 text="ביקור חיפוש",
                                                 image=search_form_icon,
                                                 width=200,
                                                 height=40,
                                                 command=self.show_visits_search_frame)
        self.search_visit_button.pack(pady=10)
        self.search_patients_button = ctk.CTkButton(self.options_frame,
                                                    text="מטופל חיפוש",
                                                    image=search_user_icon,
                                                    width=200,
                                                    height=40,
                                                    command=self.show_patients_search_frame)
        self.search_patients_button.pack(pady=10)

        self.parent_new_form_frame = ctk.CTkFrame(self.main_frame,
                                                  fg_color=color1)  # Add a parent frame inside the main window

        # To avoid multiple calls, we'll store the last resize time
        self.last_resize_time = time.time()

        self.new_form_frame = ctk.CTkFrame(
            self.parent_new_form_frame,
            fg_color="#DEEEEA",
            corner_radius=25,
            border_width=3,
            border_color="#4DBFE0"
        )
        self.current_frame = self.new_form_frame
        # Configure grid columns and rows to expand equally
        self.new_form_frame.grid_columnconfigure(0, weight=1, )  # Make column 0 expand
        self.new_form_frame.grid_columnconfigure(1, weight=1, )  # Make column 1 expand
        self.new_form_frame.grid_rowconfigure(0, weight=1, )  # Make row 0 expand
        self.new_form_frame.grid_rowconfigure(6, weight=1, )  # Make row 6 expand

        # form widgets
        self.frame_name_label = ctk.CTkLabel(
            self.new_form_frame,
            text="יצירת מטופל במערכת",
            font=hebrew_font,
            anchor="e"  # Right align the text
        )
        self.frame_name_label.grid(row=0, column=0, columnspan=2, padx=padX_size, pady=borders_widgets, sticky="we")

        self.f_name_label = ctk.CTkLabel(
            self.new_form_frame,
            text="שם פרטי",
            font=hebrew_font,
            anchor="e"  # Right align the text
        )
        self.f_name_label.grid(row=1, column=1, padx=padX_size, pady=borders_widgets, sticky=sticky_label)

        self.f_name_entry = ctk.CTkEntry(
            self.new_form_frame,
            font=hebrew_font,
            width=250,
            state="normal"

        )
        self.f_name_entry.grid(row=1, column=0, padx=padX_size, pady=borders_widgets, sticky=sticky_entry)

        # Last Name
        self.l_name_label = ctk.CTkLabel(
            self.new_form_frame,
            text="שם משפחה",
            font=hebrew_font,
            anchor="e"
        )
        self.l_name_label.grid(row=2, column=1, padx=padX_size, pady=padY_size, sticky=sticky_label)

        self.l_name_entry = ctk.CTkEntry(
            self.new_form_frame,
            font=hebrew_font,
            width=250,
            state="normal"

        )
        self.l_name_entry.grid(row=2, column=0, padx=padX_size, pady=padY_size,
                               sticky=sticky_entry)  # align the entry to the right

        # ID input
        self.id_label = ctk.CTkLabel(
            self.new_form_frame,
            text="תעודת זהות",
            font=hebrew_font,
            anchor=sticky_label
        )
        self.id_label.grid(row=3, column=1, padx=padX_size, pady=padY_size, sticky=sticky_label)

        self.id_entry = ctk.CTkEntry(
            self.new_form_frame,
            font=hebrew_font,
            width=250,
            state="normal"

        )
        self.id_entry.grid(row=3, column=0, padx=padX_size, pady=padY_size, sticky=sticky_entry)

        # Phone input
        self.phone_label = ctk.CTkLabel(
            self.new_form_frame,
            text="טלפון",
            font=hebrew_font,
            anchor="e"
        )
        self.phone_label.grid(row=4, column=1, padx=padX_size, pady=padY_size, sticky=sticky_label)

        self.phone_entry = ctk.CTkEntry(
            self.new_form_frame,
            font=hebrew_font,
            width=250,
            state="normal"

        )
        self.phone_entry.grid(row=4, column=0, padx=padX_size, pady=padY_size,
                              sticky=sticky_entry)  # align the entry to the right

        # Age Input
        self.birth_date_label = ctk.CTkLabel(
            self.new_form_frame,
            text="תאריך לידה",
            font=hebrew_font,
            anchor="e"
        )
        self.birth_date_label.grid(row=5, column=1, padx=padX_size, pady=padY_size, sticky=sticky_label)

        # Add a DateEntry widget (calendar)
        self.calendar = DateEntry(
            self.new_form_frame,
            date_pattern='dd/mm/yyyy',
            width=24,  # Increase the width to make it bigger
            background="darkblue",
            foreground="white",
            font=("Arial", 16),  # Adjust the font size to make the text inside the widget bigger
            state="normal"

        )
        self.calendar.grid(row=5, column=0, padx=padX_size, pady=padY_size, sticky=sticky_entry)
        # Submit Button
        self.create_button = ctk.CTkButton(self.new_form_frame,
                                           text="חדש מטופל צור  ",
                                           width=250,
                                           height=100,
                                           command=self.collect_data)
        self.create_button.grid(row=6, column=0, padx=padX_size, pady=(0, 30), sticky=sticky_entry)

        ###################################################################################################

        # Frame for search section
        self.search_patients_frame = ctk.CTkFrame(self.main_frame, fg_color=color1)

        #  Configure column weights properly
        self.search_patients_frame.columnconfigure(0, weight=1)  # Edit button
        self.search_patients_frame.columnconfigure(1, weight=1)  # Delete button
        self.search_patients_frame.columnconfigure(2, weight=1)  # Search button
        self.search_patients_frame.columnconfigure(3, weight=3)  # Expandable search entry
        self.search_patients_frame.columnconfigure(4, weight=2)  # Label

        #  Search Label
        self.search_label = ctk.CTkLabel(
            self.search_patients_frame,
            text="חיפוש מטופל",
            font=hebrew_font,
            anchor="center",
        )
        self.search_label.grid(row=0, column=4, padx=10, pady=10, sticky='we')

        #  Search Entry (expandable)

        self.search_patients_entry = ctk.CTkEntry(
            self.search_patients_frame,
            font=hebrew_font,
            state="normal",
            justify='right'
        )

        self.search_patients_entry.grid(row=0, column=3, padx=10, pady=10, sticky='we')  # Will expand
        self.search_patients_entry.bind("<Return>", self.search_patient_data)

        #  Search Button
        self.search_button = ctk.CTkButton(
            self.search_patients_frame,
            text="חיפוש",
            command=self.search_patient_data
        )
        self.search_button.grid(row=0, column=2, sticky='we', padx=10, pady=10)

        #  Delete Button
        self.delete_button = ctk.CTkButton(
            self.search_patients_frame,
            image=rest_icon,
            text="",
            command=self.delete_patient_data
        )
        self.delete_button.grid(row=0, column=1, sticky='we', padx=10, pady=10)

        #  TreeView Frame (Expandable)
        self.search_patients_frame.rowconfigure(1, weight=1)  # Allow row expansion

        self.patientsTreeFrame = ttk.Frame(self.search_patients_frame)
        self.patientsTreeFrame.grid(row=1, column=0, padx=10, pady=10, columnspan=5, sticky='nswe')

        self.treeScroll = ttk.Scrollbar(self.patientsTreeFrame)
        self.treeScroll.pack(side="right", fill="y")

        # Style for Treeview
        self.treeViewStyle = ttk.Style()
        self.treeViewStyle.configure("Custom.Treeview", font=("Arial ", 12))
        self.treeViewStyle.configure("Custom.Treeview.Heading", font=("Arial", 14, "bold"))

        # Columns for Treeview
        cols = ("טלפון", "גיל", "שם פרטי", "שם משפחה", "תעודה מזהה")
        self.patients_treeview = ttk.Treeview(
            self.patientsTreeFrame,
            show="headings",
            yscrollcommand=self.treeScroll.set,
            columns=cols,
            height=13,
            style="Custom.Treeview"
        )

        # Configure each column
        for col in cols:
            # Set column heading with center alignment
            self.patients_treeview.heading(col, text=col,
                                           command=lambda c=col: on_column_click(self.patients_treeview, c, False))
            # Set column width and data alignment
            self.patients_treeview.column(col, width=100, anchor="center")

        # Bind the left-click event to the open_docx function
        self.patients_treeview.bind("<Double-1>", lambda event: create_new_visit(event, self.patients_treeview))
        # Bind the Enter key press event to the open_docx function
        self.patients_treeview.bind("<Return>", lambda event: create_new_visit(event, self.patients_treeview))
        self.treeScroll.config(command=self.patients_treeview.yview)
        self.patients_treeview.pack(fill="both", expand=True)

        #  Edit Button
        self.edit_button = ctk.CTkButton(
            self.search_patients_frame,
            text="עריכה",
            command=lambda: adjust_data(self.patients_treeview)
        )
        self.edit_button.grid(row=0, column=0, sticky='we', padx=10, pady=10)

        ###################################################################################################
        self.search_visits_frame = ctk.CTkFrame(self.main_frame, fg_color=color1)  # Use ctk.CTkFrame directly

        # Configure column weights to make the layout responsive
        self.search_visits_frame.columnconfigure(0, weight=1)  # Search button
        self.search_visits_frame.columnconfigure(1, weight=1)  # Search entry
        self.search_visits_frame.columnconfigure(2, weight=3)  # Label
        self.search_visits_frame.rowconfigure(1, weight=1)  # Make treeFrame's row expandable

        self.search_label = ctk.CTkLabel(
            self.search_visits_frame,
            text="חיפוש ביקור",
            font=hebrew_font,
            anchor="center"
        )
        self.search_label.grid(row=0, column=3, padx=10, pady=5, sticky='we')

        self.search_visits_entry = ctk.CTkEntry(
            self.search_visits_frame,
            font=hebrew_font,

            justify='right'
        )

        self.search_visits = ctk.CTkEntry(
            self.search_visits_frame,
            font=hebrew_font,

            justify='right'
        )
        self.search_visits_entry.grid(row=0, column=2, padx=10, pady=10, sticky='we')
        self.search_visits_entry.bind("<Return>", self.search_visit_data)

        # Search Button
        self.search_button = ctk.CTkButton(self.search_visits_frame,
                                           text="חיפוש",
                                           width=100,
                                           command=self.search_visit_data)
        self.search_button.grid(row=0, column=1, sticky='we', padx=10, pady=10)

        self.delete_button = ctk.CTkButton(self.search_visits_frame,
                                           image=rest_icon,
                                           text="",
                                           width=50,
                                           command=self.delete_search_data)
        self.delete_button.grid(row=0, column=0, sticky='we', padx=10, pady=10)

        self.treeFrame = ttk.Frame(self.search_visits_frame)
        self.treeFrame.grid(row=1, column=0, padx=10, pady=10, columnspan=4, sticky='nswe')

        self.treeScroll = ttk.Scrollbar(self.treeFrame)
        self.treeScroll.pack(side="right", fill="y")
        self.treeViewStyle = ttk.Style()
        self.treeViewStyle.configure("Custom.Treeview",
                                     font=("Arial ", 12))
        # Configure the style for the headings with a larger font
        self.treeViewStyle.configure("Custom.Treeview.Heading",
                                     font=("Arial", 14, "bold"))  # Font for headings
        # Define columns and their headings
        sort_directions = {}  # Dictionary to track sort direction for each column
        cols = ("תאריך ביקור", "טלפון", "גיל", "שם פרטי", "שם משפחה", "תעודה מזהה")
        self.visit_treeview = ttk.Treeview(self.treeFrame,
                                           show="headings",
                                           yscrollcommand=self.treeScroll.set,
                                           columns=cols,
                                           height=13,
                                           style="Custom.Treeview")

        # Configure each column
        for col in cols:
            # Set column heading with center alignment
            self.visit_treeview.heading(col, text=col,
                                        command=lambda c=col: on_column_click(self.visit_treeview, c, sort_directions))

            # Set column width and data alignment
            self.visit_treeview.column(col, width=100, anchor="center")
        # Bind the left-click event to the open_docx function
        self.visit_treeview.bind("<Double-1>", open_word_document)
        # Bind the Enter key press event to the open_docx function
        self.visit_treeview.bind("<Return>", open_word_document)
        self.treeScroll.config(command=self.visit_treeview.yview)
        self.visit_treeview.pack(fill="both", expand=True)

        load_visit_data(self)
        load_patient_data(self)
        self.show_new_form()

    def show_new_form(self):

        self.new_form_button.configure(fg_color="#DD5746", hover_color="#C7253E")
        self.search_visit_button.configure(fg_color="#3572EF")
        self.search_patients_button.configure(fg_color="#3572EF")
        self.current_frame.pack_forget()
        self.parent_new_form_frame.pack(expand=True, fill="both", padx=20, pady=20)  # Make it responsive
        self.new_form_frame.pack(expand=True, padx=50, pady=50)
        self.current_frame = self.new_form_frame

    def show_visits_search_frame(self):

        self.search_visits_frame.pack(fill="both", expand=True)
        if self.current_frame != self.search_visits_frame:
            load_visit_data(self)
            self.search_visit_button.configure(fg_color="#DD5746", hover_color="#C7253E")
            self.new_form_button.configure(fg_color="#3572EF")
            self.search_patients_button.configure(fg_color="#3572EF")
            self.current_frame.pack_forget()
            self.parent_new_form_frame.pack_forget()
            self.current_frame = self.search_visits_frame

    def show_patients_search_frame(self):

        self.search_patients_frame.pack(fill="both", expand=True)
        if self.current_frame != self.search_patients_frame:
            load_visit_data(self)
            self.search_patients_button.configure(fg_color="#DD5746", hover_color="#C7253E")
            self.search_visit_button.configure(fg_color="#3572EF")
            self.new_form_button.configure(fg_color="#3572EF")
            self.current_frame.pack_forget()
            self.parent_new_form_frame.pack_forget()
            self.current_frame = self.search_patients_frame

    def delete_search_data(self):
        self.search_visits_entry.delete(0, tk.END)
        load_visit_data(self)

    def search_visit_data(self, event=None):
        search_term = self.search_visits_entry.get()

        # Clear existing items in visit_treeview
        for item in self.visit_treeview.get_children():
            self.visit_treeview.delete(item)

        # Get search results from database
        results = db.search_patients_visits(search_term, db_path)

        # Reinsert matching items with calculated ages
        for row in results:
            row_with_age = list(row)  # Convert the tuple to a list
            birthdate_str = row[2]  # Assuming birthdate is the 3rd column
            row_with_age[2] = calculate_age(birthdate_str)  # Replace birthdate with calculated age

            self.visit_treeview.insert('', 'end', values=row_with_age)

    def delete_patient_data(self):
        self.search_patients_entry.delete(0, tk.END)
        load_patient_data(self)

    def search_patient_data(self, event=None):
        search_term = self.search_patients_entry.get()

        # Clear existing items in visit_treeview
        for item in self.patients_treeview.get_children():
            self.patients_treeview.delete(item)

        # Get search results from database
        results = db.search_patients_data(search_term, db_path)

        # Reinsert matching items with calculated ages
        for row in results:
            row_with_age = list(row)
            birthdate_str = row[1]
            row_with_age[1] = calculate_age(birthdate_str)

            # Convert all values to string to avoid losing leading zeros
            row_with_age_str = [str(value) for value in row_with_age]

            self.patients_treeview.insert('', 'end', values=row_with_age_str)

    def collect_data(self):

        self.error_is_raised = True
        first_name = self.f_name_entry.get()
        last_name = self.l_name_entry.get()
        ID = self.id_entry.get()
        birth_date = self.calendar.get()
        phone = self.phone_entry.get()
        check_birth_date = birth_date
        if not first_name or not last_name or not birth_date or not ID or not phone:
            messagebox.showwarning("שגיאת קלט", " ! אנא מלא את כל השדות")
            return

            # Check if the birthdate format is correct
        try:

            # Assuming the expected format is 'dd/mm/yyyy'
            check_birth_date = datetime.strptime(check_birth_date, '%d/%m/%Y')

        except ValueError:
            messagebox.showerror("שגיאת קלט", "!תאריך הלידה חייב להיות בפורמט נכון: dd/mm/yyyy")
            return
        try:
            birth_date = str(birth_date)
        except ValueError:
            messagebox.showerror("שגיאת קלט", "!הגיל חייב להיות מספר")
            return

        # Get the current date in the desired format (e.g., dd-mm-yyyy)
        current_date = datetime.now().strftime('%d-%m-%Y')  # Use hyphens instead of slashes

        age = calculate_age(birth_date)
        if db.check_patient_id_exists(ID, db_path):
            messagebox.showwarning("שגיאת קלט", "המטופל כבר קיים במערכת")
        else:
            try:
                db.insert_patient_record(first_name, last_name, ID, birth_date, phone, db_path)
                docx = create_docx(first_name, last_name, ID, age, phone)
                db.insert_visit_record(ID, current_date, docx, db_path)
            except ValueError as e:
                # Catch the validation error raised by insert_patient_record and show the error message
                messagebox.showerror("Error", str(e))
                self.error_is_raised = False
            except Exception as e:
                # Catch any other errors and show a generic error message
                messagebox.showerror("Error", f"An unexpected error occurred: {str(e)}")

        if self.error_is_raised:
            # Clear all entry widgets
            self.f_name_entry.delete(0, tk.END)
            self.l_name_entry.delete(0, tk.END)
            self.id_entry.delete(0, tk.END)
            self.phone_entry.delete(0, tk.END)

        load_visit_data(self)
        load_patient_data(self)


def main():
    # Call the function to create the tables
    db.create_tables(db_path)
    root = ctk.CTk(fg_color=color1)  # create CTk window like you do with the Tk window
    PatientForm(root)
    root.mainloop()


if __name__ == "__main__":
    main()
